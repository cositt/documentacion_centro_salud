#!/usr/bin/env python3
"""Genera el documento Word completo de guías de usuario Equilibrium."""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0x99, 0x00, 0x00)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
RED_DARK = RGBColor(0x66, 0x00, 0x00)
GRAY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

BASE = os.path.dirname(os.path.abspath(__file__))

MODULES = [
    ("01_cs_resident",               "Residentes, Residencias, Habitaciones y Trabajadores"),
    ("02_cs_purse_pocket",           "Monedero y Fondos de Terceros"),
    ("03_cs_admission",              "Proceso de Admisión Guiado"),
    ("04_cs_medical_care",           "Atención Médica"),
    ("05_cs_cima",                   "Catálogo de Medicamentos CIMA"),
    ("06_cs_psychology",             "Psicología y Llamadas Familiares"),
    ("07_cs_monitoring",             "Monitorización Diaria"),
    ("08_cs_patient_followup_forms", "Motor de Formularios Dinámicos"),
    ("09_cs_piai",                   "Plan Individualizado de Atención Integral"),
    ("10_cs_documents",              "Gestión de Documentos Clínicos"),
    ("11_cs_dni_ocr",                "OCR de DNI/NIE"),
    ("12_cs_residential_stock",      "Inventario Residencial"),
]


# ── helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def page_break(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def red_rule(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 65)
    run.font.color.rgb = RED
    run.font.size = Pt(8)


# ── estilos ───────────────────────────────────────────────────────────────────

def setup_styles(doc: Document):
    s = doc.styles

    n = s["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(11)
    n.font.color.rgb = BLACK
    n.paragraph_format.space_after = Pt(6)
    n.paragraph_format.line_spacing = Pt(14)

    h1 = s["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = RED
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = False

    h2 = s["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(15)
    h2.font.bold = True
    h2.font.color.rgb = RED
    h2.paragraph_format.space_before = Pt(16)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = s["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RED_DARK
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    h4 = s["Heading 4"]
    h4.font.name = "Calibri"
    h4.font.size = Pt(11)
    h4.font.bold = True
    h4.font.italic = True
    h4.font.color.rgb = BLACK
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.space_after = Pt(2)

    lb = s["List Bullet"]
    lb.font.name = "Calibri"
    lb.font.size = Pt(11)
    lb.font.color.rgb = BLACK
    lb.paragraph_format.space_after = Pt(3)

    try:
        cs = s.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    except Exception:
        cs = s["Normal"]
    cs.font.name = "Courier New"
    cs.font.size = Pt(9)
    cs.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    cs.paragraph_format.space_before = Pt(4)
    cs.paragraph_format.space_after = Pt(4)
    cs.paragraph_format.left_indent = Cm(0.6)


# ── portada ───────────────────────────────────────────────────────────────────

def add_cover(doc: Document):
    for _ in range(7):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 38)
    run.font.color.rgb = RED
    run.font.size = Pt(16)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SISTEMA EQUILIBRIUM")
    run.font.name = "Calibri"
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RED

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Guías de Usuario Completas")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = BLACK

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Centro Residencial Equilibrium")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = GRAY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Odoo 19.0 Enterprise · 12 Módulos Custom cs_")
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.color.rgb = GRAY

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 38)
    run.font.color.rgb = RED
    run.font.size = Pt(16)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Módulos documentados")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RED

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, txt in enumerate(["#", "Módulo", "Área funcional"]):
        hdr[i].text = txt
        r = hdr[i].paragraphs[0].runs[0]
        r.font.bold = True
        r.font.color.rgb = WHITE
        r.font.name = "Calibri"
        set_cell_bg(hdr[i], "990000")

    for idx, (folder, title) in enumerate(MODULES, 1):
        parts = folder.split("_", 1)
        mod_name = "cs_" + parts[1] if len(parts) > 1 else folder
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = mod_name
        row[2].text = title
        for cell in row:
            for run in cell.paragraphs[0].runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10)
        if idx % 2 == 0:
            for cell in row:
                set_cell_bg(cell, "F5E6E6")

    for row in table.rows:
        widths = [Cm(1.2), Cm(5.5), Cm(9.5)]
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]

    page_break(doc)


# ── índice manual ─────────────────────────────────────────────────────────────

def collect_headings():
    """Recorre todos los guia.md y extrae (nivel, texto, módulo)."""
    entries = []
    for folder, _ in MODULES:
        path = os.path.join(BASE, folder, "guia.md")
        if not os.path.exists(path):
            continue
        with open(path, encoding="utf-8") as f:
            for line in f:
                line = line.rstrip()
                m4 = re.match(r"^#### (.+)", line)
                m3 = re.match(r"^### (.+)", line)
                m2 = re.match(r"^## (.+)", line)
                m1 = re.match(r"^# (.+)", line)
                if m1:
                    entries.append((1, m1.group(1), folder))
                elif m2:
                    entries.append((2, m2.group(1), folder))
                elif m3:
                    entries.append((3, m3.group(1), folder))
    return entries


def add_manual_toc(doc: Document):
    p = doc.add_paragraph("ÍNDICE DE CONTENIDOS", style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    headings = collect_headings()

    indent_map = {1: Cm(0), 2: Cm(0.6), 3: Cm(1.4)}
    size_map   = {1: Pt(12), 2: Pt(11), 3: Pt(10)}
    bold_map   = {1: True, 2: False, 3: False}
    color_map  = {1: RED, 2: BLACK, 3: GRAY}

    for level, text, _folder in headings:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2) if level > 1 else Pt(8)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = indent_map[level]

        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = size_map[level]
        run.font.bold = bold_map[level]
        run.font.color.rgb = color_map[level]

    page_break(doc)


# ── parser markdown → docx ────────────────────────────────────────────────────

def parse_markdown(doc: Document, text: str):
    lines = text.split("\n")
    in_code = False
    code_buf = []
    i = 0

    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            if in_code:
                if code_buf:
                    p = doc.add_paragraph("\n".join(code_buf), style="CodeBlock")
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        stripped = line.strip()

        if line.startswith("#### "):
            doc.add_paragraph(line[5:].strip(), style="Heading 4")
        elif line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 3")
        elif line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 2")
        elif line.startswith("# "):
            doc.add_paragraph(line[2:].strip(), style="Heading 1")
        elif stripped in ("---", "***", "___"):
            red_rule(doc)
        elif re.match(r"^    [-*+] ", line) or re.match(r"^  [-*+] ", line):
            content = re.sub(r"^\s+[-*+] ", "", line).strip()
            p = doc.add_paragraph(style="List Bullet 2")
            add_inline(p, content)
        elif re.match(r"^[-*+] ", line):
            content = line[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, content)
        elif re.match(r"^\d+\. ", line):
            content = re.sub(r"^\d+\. ", "", line).strip()
            p = doc.add_paragraph(style="List Number")
            add_inline(p, content)
        elif line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[-| :]+\|", lines[i + 1]):
            rows = []
            j = i
            while j < len(lines) and lines[j].startswith("|"):
                cells = [c.strip() for c in lines[j].split("|")]
                cells = [c for c in cells if c]
                if cells and not all(re.match(r"^[-: ]+$", c) for c in cells):
                    rows.append(cells)
                j += 1
            i = j
            if rows:
                cols = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=cols)
                tbl.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci in range(cols):
                        cell_text = row[ci] if ci < len(row) else ""
                        cell = tbl.rows[ri].cells[ci]
                        p = cell.paragraphs[0]
                        p.clear()
                        add_inline(p, cell_text)
                        for run in p.runs:
                            run.font.name = "Calibri"
                            run.font.size = Pt(10)
                        if ri == 0:
                            for run in p.runs:
                                run.font.bold = True
                                run.font.color.rgb = WHITE
                            set_cell_bg(cell, "990000")
                        elif ri % 2 == 0:
                            set_cell_bg(cell, "F5E6E6")
                doc.add_paragraph()
            continue
        elif stripped == "":
            pass
        else:
            p = doc.add_paragraph()
            add_inline(p, line)

        i += 1


def add_inline(paragraph, text: str):
    i = 0
    while i < len(text):
        # bold+italic
        if text[i:i+3] == "***":
            end = text.find("***", i + 3)
            if end != -1:
                run = paragraph.add_run(text[i+3:end])
                run.bold = True
                run.italic = True
                run.font.color.rgb = BLACK
                i = end + 3
                continue
        # bold
        if text[i:i+2] == "**":
            end = text.find("**", i + 2)
            if end != -1:
                run = paragraph.add_run(text[i+2:end])
                run.bold = True
                run.font.color.rgb = BLACK
                i = end + 2
                continue
        # italic
        if text[i] == "*":
            end = text.find("*", i + 1)
            if end != -1:
                run = paragraph.add_run(text[i+1:end])
                run.italic = True
                run.font.color.rgb = BLACK
                i = end + 1
                continue
        # code
        if text[i] == "`":
            end = text.find("`", i + 1)
            if end != -1:
                run = paragraph.add_run(text[i+1:end])
                run.font.name = "Courier New"
                run.font.size = Pt(10)
                run.font.color.rgb = RED_DARK
                i = end + 1
                continue
        # texto plano: avanzar hasta próximo marcador (o fin)
        next_marker = len(text)
        for marker in ("***", "**", "*", "`"):
            pos = text.find(marker, i + 1)  # +1 evita quedarse atascado en marcador sin par
            if pos != -1 and pos < next_marker:
                next_marker = pos
        run = paragraph.add_run(text[i:next_marker])
        run.font.color.rgb = BLACK
        i = next_marker


# ── main ──────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    setup_styles(doc)
    add_cover(doc)
    add_manual_toc(doc)

    for folder, _ in MODULES:
        path = os.path.join(BASE, folder, "guia.md")
        if not os.path.exists(path):
            continue
        with open(path, encoding="utf-8") as f:
            content = f.read()
        parse_markdown(doc, content)
        page_break(doc)

    output = os.path.join(BASE, "Equilibrium_Guias_Usuario.docx")
    doc.save(output)
    print(f"✓ Guardado: {output}")


if __name__ == "__main__":
    main()
